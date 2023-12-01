import streamlit as st
import openai
import os
import json
from langchain.utilities.dalle_image_generator import DallEAPIWrapper
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.llms import OpenAI
from dotenv import load_dotenv
from docx import Document
document = Document()

# Load environment variables
load_dotenv()

# Configure OpenAI API
openai.api_key = os.getenv('OPENAI_API_KEY')

# Chat UI title
st.header("Blog Generator")

question_generator_prompt = """You are a efficient Blog writer with 20 years of experience.
    Your job is to create a compelling blog post for the details given by the user.
    The blog must be creative and should limit to 1000 words.
    The language to be used is English.
    The result must contain a blog Heading, the content paragraphs as a single text file  
    The result should be return in a json file with keys as heading, content and the corresponding values. 
    For example:
    Answer: 
    {"heading": "heading of the blog",
    "content" : "content of the blog",}
    """
    
#Initialising llm for image generation prompt
llm = OpenAI(temperature=0.9)
prompt = PromptTemplate(
    input_variables=["image_desc"],
    template="Generate a detailed prompt to generate an image based on the following description: {image_desc}",
)
chain = LLMChain(llm=llm, prompt=prompt)

def blog_generator(keyword):
    
    completion = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",
    temperature = 0.5,
    max_tokens=1500,
    messages=[
        {"role": "system", "content": question_generator_prompt},
        {"role": "user", "content": "Create a blog based on the instructions given below.\nTopic: "+keyword}
    ]
    )
    
    return (completion.choices[0].message['content'])

def image_generator(image_prompt):
    image_url = DallEAPIWrapper().run(chain.run(image_prompt))
    return(image_url)

        
with st.sidebar:
    #uploaded_file = st.file_uploader("Please upload your file",type=None)# accept_multiple_files=True, 
    topic_name=st.text_input('Details for the blog')
    #topic_name = st.text_area("Details for the blog")
    if topic_name:
        st.write(topic_name)
    analyze = st.button("Generate blog")
    #save = st.button("Save the blog")
    
if topic_name and analyze:
    #st.write("Generating the blog")
    
    with st.spinner('Generating the blog'):
        result = json.loads(blog_generator(topic_name))
        #st.write(result["heading"])
        url = image_generator(topic_name)
    
    #print(url)
    #print("Prompt: ",result["prompt"])
    heading = result["heading"]
    st.write(f"**{heading}**")
    st.image(url,width=400)
    st.write(result["content"])
    #st.write(result["prompt"])
    #if save:
    document.add_heading(result["heading"]).bold = True
    paragraph = document.add_paragraph()
    paragraph.add_run("\n")
    paragraph.add_run(result["content"])
    document.save('New_blog.docx')
    
else:
    st.write("Write a topic to generate Blog")    
